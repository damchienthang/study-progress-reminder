import docx
import sys

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.replace('\n', ' ').strip())
                full_text.append(" | ".join(row_data))
        print('\n'.join(full_text))
    except Exception as e:
        print(f"Error: {e}")

if __name__ == '__main__':
    read_docx(sys.argv[1])
